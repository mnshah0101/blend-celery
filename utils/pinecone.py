import json
import boto3
from urllib.parse import urlparse
import io
import os
import re
from langchain_openai import ChatOpenAI
from langchain_core.prompts import PromptTemplate

import requests as requests

import boto3
from urllib.parse import urlparse
from io import BytesIO
import docx
from langchain.docstore.document import Document
from langchain_text_splitters.character import RecursiveCharacterTextSplitter

from langchain_community.document_loaders.pdf import PyPDFLoader
from langchain_openai import OpenAIEmbeddings

from langchain_pinecone import PineconeVectorStore
from collections import defaultdict


import dotenv

dotenv.load_dotenv()

openai_api_key = os.getenv('OPENAI_API_KEY')
pinecone_api_key = os.getenv('PINECONE_API_KEY')
go_server_url = os.getenv('GO_SERVER_URL')
bucket = os.getenv('BUCKET_NAME')


def read_s3_pdf_from_url(file_url):
    try:
        print(f"Reading PDF from URL: {file_url}")
        loader = PyPDFLoader(file_url)
        data = loader.load()

        file_url = data[0].metadata['source']
        doc_id = getDocumentIdByUrl(file_url)

        for i in range(len(data)):
            data[i].metadata['id'] = doc_id

        print('this id tha data')
        print(data)
        print("PDF data loaded successfully")
        return data
    except Exception as e:
        print(f"Error reading PDF from URL: {e}")
        return None


def read_s3_text_file_from_url(s3_url):
    print(f"Reading text file from URL: {s3_url}")
    parsed_url = urlparse(s3_url)
    bucket_name = bucket
    file_key = parsed_url.path.lstrip('/')
    file_key = re.sub(r'\+', ' ', file_key)
    s3_client = boto3.client('s3')
    try:
        response = s3_client.get_object(Bucket=bucket_name, Key=file_key)
        content = response['Body'].read().decode('utf-8')
        print(f"Text file content loaded successfully from {file_key}")

        # Split content by pages based on a delimiter, e.g., '\f' for form feed
        pages = content.split('\f')
        print("Text file content split into pages")

        documents = []
        for num, text in enumerate(pages):
            print(f"Processing page {num} from text file {file_key}")

            doc_id = getDocumentIdByUrl(s3_url)

            doc = Document(
                metadata={"page": num, "source": file_key, "id": doc_id}, page_content=text)
            documents.append(doc)
            print(f"Page {num} added as Document")

        return documents
    except Exception as e:
        print(f"Error reading text S3 file: {e}")
        return None


def read_s3_docx_from_url(s3_url):
    print(f"Reading DOCX file from URL: {s3_url}")
    parsed_url = urlparse(s3_url)
    bucket_name = bucket
    file_key = parsed_url.path.lstrip('/')
    file_key = re.sub(r'\+', ' ', file_key)
    s3_client = boto3.client('s3')

    try:
        response = s3_client.get_object(Bucket=bucket_name, Key=file_key)
        file_stream = BytesIO(response['Body'].read())
        doc = docx.Document(file_stream)
        print(f"DOCX file content loaded successfully from {file_key}")

        full_text = []
        current_page = []

        for paragraph in doc.paragraphs:
            current_page.append(paragraph.text)
            # Check for page breaks, which are typically added as separate paragraphs with '\f'
            if paragraph.text == '\f':
                full_text.append('\n'.join(current_page))
                current_page = []

        # Add the last page
        if current_page:
            full_text.append('\n'.join(current_page))
        documents = []
        for num, text in enumerate(full_text):

            doc_id = getDocumentIdByUrl(s3_url)

            doc = Document(
                metadata={"page": num, "source": s3_url, "id": doc_id}, page_content=text)
            documents.append(doc)
            print(f"Page {num} added as Document")

        return documents
    except Exception as e:
        print(f"Error reading docx S3 file: {e}")
        return None

    # https://avaloncasesbucket.s3.amazonaws.com/ff8f4f2a60e94647f34ee5167728ca74/20240524152459Resume_.docx
    # https://avaloncasesbucket.s3.amazonaws.com/ff8f4f2a60e94647f34ee5167728ca74/20240524151507Resume_.docx


def create_embeddings(namespaces, docs):
    try:
        print("Creating embeddings")
        print("length of docs" + str(len(docs)))
        print("length of namespaces" + str(len(namespaces)))
        embeddings = OpenAIEmbeddings(
            openai_api_key=openai_api_key, model="text-embedding-3-large")

        tracker = {}

        for i in range(len(docs)):
            if namespaces[i] not in tracker:
                tracker[namespaces[i]] = []
            tracker[namespaces[i]].append(docs[i])

        for key in tracker:
            print("Creating Pinecone Vector Store")
            vectorstore = PineconeVectorStore(
                index_name="avalondocbucket", embedding=embeddings, pinecone_api_key=pinecone_api_key, namespace=key)

            docs_to_upload = tracker[key]

            for doc in docs_to_upload:
                if (doc.metadata['page'] != 0):
                    continue
                relevancy = create_relevancy_score(doc, key)
                print("Relevancy Score: " + str(relevancy))
                first_doc = doc
                first_doc_url = first_doc.metadata['source']
                try:
                    response = requests.post(
                        f"{go_server_url}/updateRelevancyByFileUrl", json={"file_url": first_doc_url, "relevancy": relevancy})
                    print(response.json())
                except Exception as e:
                    print(f"Error updating relevancy score: {e}")

            print("Creating Recursive Character Text Splitter")
            splitter = RecursiveCharacterTextSplitter()
            split_docs = splitter.split_documents(docs_to_upload)
            print(f"Document split into {len(split_docs)} chunks")
            vectorstore.add_documents(split_docs)
            print(
                f"Chunks added to Pinecone index")

        return {"status": "uploaded", "message": "Documents uploaded"}
    except Exception as e:
        print(f"Error creating embeddings: {e}")
        return {"status": "error", "message": f"Error Message: {e}"}


def read_s3_generic_file_from_url(s3_url):
    print(f"Reading generic file from URL: {s3_url}")
    parsed_url = urlparse(s3_url)
    bucket_name = bucket
    file_key = parsed_url.path.lstrip('/')
    # change + to space
    file_key = re.sub(r'\+', ' ', file_key)

    print(f"Reading file from bucket: {bucket_name}, key: {file_key}")
    s3_client = boto3.client('s3')

    try:
        response = s3_client.get_object(Bucket=bucket_name, Key=file_key)

        # Use file extension to determine how to process the file
        if file_key.endswith('.txt'):
            return read_s3_text_file_from_url(s3_url)
        elif file_key.endswith('.docx'):
            return read_s3_docx_from_url(s3_url)
        elif file_key.endswith('.pdf'):
            return read_s3_pdf_from_url(s3_url)
        else:
            # Fallback for unsupported file types
            print(f"Unsupported file type for: {file_key}")
            return None
    except Exception as e:
        print(f"Error reading S3 file from generic: {e}")
        return None


def read_documents(s3_urls):
    print(f"Reading documents from URLs: {s3_urls}")
    documents = []
    case_ids = []

    for s3_url in s3_urls:
        s3_id = s3_url.strip().split('/')[3]
        print(f"Case ID: {s3_id}")

        doc = read_s3_generic_file_from_url(s3_url)
        if doc:
            append_id = [s3_id] * len(doc)
            case_ids.extend(append_id)
            documents.extend(doc)
            print(f"Document from {s3_url} added to documents list")
        else:
            print(f"Error reading document from in read_documents {s3_url}")
            continue
    return case_ids, documents


def upload_documents(s3_urls, ids, file_names):
    try:

        case_ids, documents = read_documents(s3_urls)

        response = create_embeddings(case_ids, documents)
        for i in range(len(ids)):
            updateDynamoStatus(ids[i], file_names[i])
        print("All documents uploaded successfully")
        return {"status": "uploaded", "message": "Documents uploaded successfully"}
    except Exception as e:
        print(f"Error uploading documents: {e}")
        return {"status": "error", "message": f"Error Message {e}"}


def getProcessingDocs():
    print("Fetching documents from DynamoDB")
    dynamodb = boto3.resource('dynamodb')

    table = dynamodb.Table('AvalonDocuments')

    response = table.scan(
        FilterExpression=boto3.dynamodb.conditions.Attr('stored').eq(False)
    )

    s3_urls = [doc['file_url']
               for doc in response.get('Items', []) if 'file_url' in doc]
    if not s3_urls:
        return [], [], []

    ids = [doc['_id'] for doc in response.get('Items', []) if '_id' in doc]
    file_names = [doc['file_name']
                  for doc in response.get('Items', []) if 'file_name' in doc]
    return s3_urls, ids, file_names


def updateDynamoStatus(_id, file_name):
    dynamodb = boto3.resource('dynamodb')
    table = dynamodb.Table('AvalonDocuments')
    table.update_item(
        Key={
            '_id': _id,
        },
        UpdateExpression="SET #s = :s",
        ExpressionAttributeValues={
            ':s': True
        },
        ExpressionAttributeNames={
            '#s': 'stored'
        },
        ReturnValues="UPDATED_NEW"
    )


def lambda_handler(event, context):
    print("Lambda handler triggered")
    index_name, s3_urls, ids, file_names = getProcessingDocs()
    if not s3_urls:
        print("No documents to process")
        return {
            'statusCode': 200,
            'body': 'No documents to process'
        }

    response = upload_documents(index_name, s3_urls, ids, file_names)
    if response['status'] == 'uploaded':
        print("Documents uploaded successfully in lambda handler")
        return {
            'statusCode': 200,
            'body': 'Uploaded Successfully'
        }
    print("Error during document upload in lambda handler")
    return {
        'statusCode': 500,
        'body': 'Error during upload'
    }


def create_relevancy_score(doc, case_id):
    try:

        page_content = doc.page_content
        print('case_id')
        print(case_id)

        new_case = requests.post(
            f"{go_server_url}/getCase", json={"_id": case_id}).json()['object']

        case_info = new_case['case_info']

        template = """

        You are an AI Legal Assistant that takes in the text of a legal document: {page_content} and assigns a score
        on a scale from 0-100 based on their relevancy to the overall facts of the case presented: {case_info}. Only return an integer, and nothing but an integer

        Example response: "80".

        """

        prompt = PromptTemplate.from_template(template)

        print("this it he prompt")
        print(prompt)

        llm = ChatOpenAI(model="gpt-4o")

        llm_chain = prompt | llm

        relevancy_score = llm_chain.invoke(
            {"page_content": page_content, "case_info": case_info})

        relevancy_score = relevancy_score.content

        print("Relevancy Score: " + str(relevancy_score))

        # get rid of any non-numeric characters
        relevancy_score = re.sub(r'\D', '', relevancy_score)

        return int(relevancy_score)

    except Exception as e:
        print(f"Error creating relevancy score: {e}")
        return 0


def getDocumentIdByUrl(file_url):
    try:
        response = requests.post(
            f"{go_server_url}/getDocumentIdByUrl", json={"file_url": file_url})
        print(response.json())
        return response.json()['object']
    except Exception as e:
        print(f"Error getting document ID by URL: {e}")
        return ''


def handleUpdate():
    try:
        s3_urls, ids, file_names = getProcessingDocs()

        for i in range(len(ids)):
            response = upload_documents(
                [s3_urls[i]], [ids[i]], [file_names[i]])
            if response['status'] == 'uploaded':
                print("Document uploaded successfully")
            else:
                print("Error uploading document")

        return {"status": 200, "data": "Documents uploaded"}
    except Exception as e:
        print(e)
        return {"status": 500, "message": str(e)}
